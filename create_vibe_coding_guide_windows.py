"""
Script to create the Vibe Coding Setup Guide Word document - WINDOWS VERSION
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_code_block(doc, code_text):
    """Add a code block with monospace formatting"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading)


def add_tip_box(doc, tip_text):
    """Add a tip/note callout"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("TIP: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)
    p.add_run(tip_text)


def add_warning_box(doc, warning_text):
    """Add a warning callout"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("WARNING: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run(warning_text)


def add_success_box(doc, success_text):
    """Add a success/verification callout"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("SUCCESS LOOKS LIKE: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 150)
    p.add_run(success_text)


def add_screenshot_placeholder(doc, description):
    """Add a placeholder for screenshots"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[SCREENSHOT: {description}]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_numbered_list(doc, items):
    """Add numbered list items"""
    for item in items:
        doc.add_paragraph(item, style='List Number')


def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('Vibe Coding Setup Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('A Step-by-Step Guide to AI-Assisted Development with Claude CLI')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    version = doc.add_paragraph('For Windows Users | No Coding Experience Required')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # =========================================================================
    # BEFORE YOU BEGIN
    # =========================================================================
    doc.add_heading('Before You Begin', level=1)

    doc.add_heading('What You Need', level=2)

    prereqs = [
        'A Windows 10 or Windows 11 computer',
        'Admin access to your PC (ability to install software)',
        'An internet connection',
        'About 2 hours total (split across multiple sessions)',
        'An Anthropic account for Claude (we\'ll create this in Phase 2)'
    ]
    for item in prereqs:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('What You\'ll Accomplish', level=2)

    doc.add_paragraph(
        'By the end of this guide, you\'ll be able to describe a task in plain English and '
        'watch AI write the code for you. Want to analyze a spreadsheet? Automate a repetitive task? '
        'Generate reports? You\'ll tell Claude what you need, and it will build it.'
    )

    doc.add_heading('A Note About the Terminal', level=2)

    doc.add_paragraph(
        'Much of this guide involves typing commands into something called the "Terminal" (also called '
        '"Command Prompt" or "PowerShell" on Windows). The Terminal is just a text-based way to talk to '
        'your computer. Instead of clicking buttons, you type commands. It looks intimidating at first, '
        'but you\'ll only need to learn a handful of commands, and we\'ll walk you through each one.'
    )

    doc.add_paragraph(
        'Think of the Terminal like texting your computer: you type a message, press Enter, '
        'and your computer responds.'
    )

    doc.add_page_break()

    # =========================================================================
    # OVERVIEW SECTION
    # =========================================================================
    doc.add_heading('Overview: What is Vibe Coding?', level=1)

    doc.add_paragraph(
        '"Vibe coding" is a new approach to building software where you describe what you want in plain English, '
        'and an AI assistant (Claude) writes the code for you. You don\'t need to be a programmer—you just need '
        'to clearly describe your goal, review what Claude creates, and guide it with feedback.'
    )

    doc.add_paragraph(
        'It\'s called "vibe" coding because you\'re working at a higher level—focusing on WHAT you want, '
        'not HOW to code it. You set the direction, Claude handles the technical details.'
    )

    doc.add_heading('The Four Components You\'ll Set Up', level=2)

    add_screenshot_placeholder(doc, "Diagram showing how VS Code, Python, Claude CLI, and GitHub connect")

    components = [
        ('VS Code', 'Your workspace. A free code editor from Microsoft. Think of it as a "smart notepad" designed for code. '
         'It shows your files on the left, your code in the middle, and has a built-in Terminal at the bottom '
         'where you\'ll type commands.'),
        ('Python', 'The language. A programming language that Claude will write code in. You don\'t need to learn Python—'
         'Claude handles that—but Python needs to be installed so the code can actually run.'),
        ('Claude CLI', 'Your AI partner. The command-line interface for Claude AI. This is how you\'ll "talk" to Claude '
         'directly from VS Code\'s Terminal. You describe what you want, Claude writes and edits files for you.'),
        ('GitHub', 'Your safety net. A cloud service that stores your code and tracks every change. Think of it as '
         '"Google Docs version history" for code. If something breaks, you can always go back.')
    ]

    for name, desc in components:
        p = doc.add_paragraph()
        run = p.add_run(f"{name} — ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading('How a Vibe Coding Session Works', level=2)

    doc.add_paragraph('Once everything is set up, here\'s what a typical 30-minute session looks like:')

    workflow = [
        'You open VS Code and see your project files in the left sidebar',
        'You open the Terminal (a panel at the bottom) by pressing Ctrl+` (backtick key)',
        'You type "claude" and press Enter to start talking to the AI',
        'You describe what you want in plain English: "Create a script that reads my Excel file and summarizes sales by region"',
        'Claude writes the code, creates files, and explains what it did. You\'ll see new files appear in the sidebar.',
        'You run the code to see if it works (Claude often does this for you)',
        'If it doesn\'t work perfectly, you tell Claude what to fix: "The date format is wrong, use MM/DD/YYYY"',
        'When you\'re happy, you save your progress to GitHub so you never lose your work'
    ]
    add_numbered_list(doc, workflow)

    doc.add_page_break()

    # =========================================================================
    # PHASE 1
    # =========================================================================
    doc.add_heading('Phase 1: Foundation Setup', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Time estimate: 30-45 minutes')
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('Goal: ')
    run.bold = True
    p.add_run('Install the basic tools and get comfortable with VS Code\'s layout.')

    doc.add_paragraph(
        'After completing Phase 1, you\'ll have VS Code running and understand where everything is on screen. '
        'You can stop here and come back later for Phase 2.'
    )

    # Step 1.1
    doc.add_heading('Step 1.1: Install VS Code', level=2)

    steps_1_1 = [
        'Open your browser and go to: https://code.visualstudio.com/',
        'Click the big blue "Download for Windows" button',
        'Run the downloaded installer (VSCodeUserSetup-x64-x.xx.x.exe)',
        'Accept the license agreement and click Next',
        'IMPORTANT: On the "Select Additional Tasks" screen, check ALL the boxes, especially "Add to PATH"',
        'Click Install, then Finish when done',
        'Open VS Code from the Start Menu or desktop shortcut'
    ]
    add_numbered_list(doc, steps_1_1)

    add_success_box(doc, 'A window opens with "Welcome" tab. You\'ll see a dark (or light) interface with icons on the left side.')

    add_screenshot_placeholder(doc, "VS Code welcome screen on Windows")

    # VS Code Orientation
    doc.add_heading('Step 1.2: Learn the VS Code Layout', level=2)

    doc.add_paragraph(
        'Before we install anything else, let\'s understand what you\'re looking at. '
        'This is worth spending 5 minutes on—it\'ll make everything else easier.'
    )

    add_screenshot_placeholder(doc, "VS Code with labeled areas: Sidebar, Editor, Terminal, Activity Bar, Status Bar")

    layout_items = [
        ('Activity Bar (far left, vertical icons)',
         'These icons switch between different views. The top icon (two overlapping files) is "Explorer"—this is the one you\'ll use 99% of the time.'),
        ('Sidebar (left panel)',
         'When Explorer is selected, this shows your folders and files, like File Explorer. Click a file to open it. Right-click to create new files/folders.'),
        ('Editor Area (large center area)',
         'This is where you view and edit files. Files open in tabs at the top, just like browser tabs. You can have multiple files open.'),
        ('Terminal (bottom panel)',
         'THIS IS THE IMPORTANT ONE. You\'ll type commands here. If you don\'t see it, press Ctrl+` (backtick, the key left of 1). We\'ll use this constantly.'),
        ('Status Bar (very bottom strip)',
         'Shows info like which line you\'re on and which Python version is active. You can mostly ignore this for now.')
    ]

    for name, desc in layout_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading('Try it now:', level=3)

    try_it = [
        'Press Ctrl+` (hold Control, tap the backtick key). The Terminal should appear at the bottom.',
        'Press Ctrl+` again. The Terminal hides.',
        'Press Ctrl+` once more to bring it back.',
        'Click in the Terminal area. You should see a blinking cursor and a prompt like "PS C:\\Users\\YourName>"'
    ]
    add_numbered_list(doc, try_it)

    add_tip_box(doc, 'The Terminal is your command center. Ctrl+` is the most important keyboard shortcut you\'ll learn.')

    # Step 1.3
    doc.add_heading('Step 1.3: Install Python', level=2)

    doc.add_paragraph(
        'Python is the programming language Claude will write code in. Let\'s install it from the official source.'
    )

    steps_python = [
        'Go to: https://www.python.org/downloads/',
        'Click the big yellow "Download Python 3.x.x" button',
        'Run the downloaded installer',
        'CRITICAL: Check the box that says "Add Python to PATH" at the bottom of the installer!',
        'Click "Install Now"',
        'Wait for installation to complete, then click Close'
    ]
    add_numbered_list(doc, steps_python)

    add_warning_box(doc, 'If you skip "Add Python to PATH," Python won\'t work from the Terminal. If you missed it, uninstall Python and reinstall with that box checked.')

    doc.add_paragraph('Close and reopen VS Code, then verify Python installed correctly. In the Terminal:')
    add_code_block(doc, 'python --version')

    add_success_box(doc, 'You see "Python 3.12.x" or "Python 3.13.x" (the exact number may vary)')

    add_tip_box(doc, 'On Windows, use "python" (not "python3" like on Mac).')

    # Step 1.4
    doc.add_heading('Step 1.4: Install Git', level=2)

    doc.add_paragraph('Git is version control software—it tracks changes to your code.')

    steps_git = [
        'Go to: https://git-scm.com/download/win',
        'The download should start automatically. If not, click "Click here to download"',
        'Run the installer',
        'Click Next through the screens, accepting defaults (there are many screens—defaults are fine)',
        'On "Choosing the default editor," you can select "Use Visual Studio Code as Git\'s default editor"',
        'Continue clicking Next, then Install, then Finish'
    ]
    add_numbered_list(doc, steps_git)

    doc.add_paragraph('Close and reopen VS Code (important!), then verify Git installed:')
    add_code_block(doc, 'git --version')

    add_success_box(doc, 'You see "git version 2.x.x"')

    doc.add_paragraph('Configure Git with your name and email. This labels your changes:')
    add_code_block(doc, 'git config --global user.name "Your Name"')
    doc.add_paragraph('(Replace "Your Name" with your actual name, keep the quotes)')

    add_code_block(doc, 'git config --global user.email "your.email@example.com"')
    doc.add_paragraph('(Replace with your actual email, keep the quotes)')

    add_tip_box(doc, 'Use the same email you\'ll use for your GitHub account in Phase 4.')

    # Step 1.5
    doc.add_heading('Step 1.5: Create Your Projects Folder', level=2)

    doc.add_paragraph('Let\'s create a dedicated folder for all your coding projects. In the Terminal:')
    add_code_block(doc, 'mkdir C:\\Users\\%USERNAME%\\Documents\\projects')

    doc.add_paragraph(
        'This creates a folder called "projects" inside your Documents folder.'
    )

    doc.add_paragraph('Now open this folder in VS Code:')
    add_code_block(doc, 'code C:\\Users\\%USERNAME%\\Documents\\projects')

    add_success_box(doc, 'A new VS Code window opens. You\'ll see "PROJECTS" in the Explorer sidebar on the left (it will be empty).')

    add_tip_box(doc, 'From now on, you\'ll do all your work inside this projects folder. Each project gets its own subfolder.')

    # Phase 1 Checkpoint
    doc.add_heading('Phase 1 Checkpoint', level=2)

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run('Before moving on, verify each item:')
    run.bold = True

    checklist_1 = [
        'VS Code opens and you know how to show/hide the Terminal (Ctrl+`)',
        'Typing "python --version" shows "Python 3.x.x"',
        'Typing "git --version" shows "git version x.x.x"',
        'You have VS Code open with your projects folder showing in the sidebar'
    ]
    for item in checklist_1:
        doc.add_paragraph("☐ " + item)

    p = doc.add_paragraph()
    run = p.add_run('Great stopping point! ')
    run.bold = True
    p.add_run('Take a break here if needed. Phase 2 is where Claude enters the picture, and it\'s the exciting part!')

    doc.add_page_break()

    # =========================================================================
    # PHASE 2
    # =========================================================================
    doc.add_heading('Phase 2: Claude CLI Setup', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Time estimate: 15-20 minutes')
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('Goal: ')
    run.bold = True
    p.add_run('Install Claude CLI and have your first AI-assisted coding conversation.')

    # Step 2.1
    doc.add_heading('Step 2.1: Install Node.js', level=2)

    doc.add_paragraph(
        'Claude CLI is built with Node.js (a programming platform), so we need to install that first.'
    )

    steps_node = [
        'Go to: https://nodejs.org/',
        'Click the LTS (Long Term Support) version button—it\'s the recommended one',
        'Run the downloaded installer',
        'Click Next through the screens, accepting defaults',
        'Click Install, then Finish'
    ]
    add_numbered_list(doc, steps_node)

    doc.add_paragraph('Close and reopen VS Code, then verify in the Terminal:')
    add_code_block(doc, 'node --version')

    add_success_box(doc, 'You see "v20.x.x" or "v22.x.x" (or similar)')

    # Step 2.2
    doc.add_heading('Step 2.2: Install Claude CLI', level=2)

    doc.add_paragraph('Now install Claude CLI globally (so you can use it from any folder):')
    add_code_block(doc, 'npm install -g @anthropic-ai/claude-code')

    doc.add_paragraph('This takes a minute or two. When it finishes, verify:')
    add_code_block(doc, 'claude --version')

    add_success_box(doc, 'You see a version number like "1.x.x"')

    # Step 2.3
    doc.add_heading('Step 2.3: Create an Anthropic Account', level=2)

    doc.add_paragraph(
        'Claude CLI requires an Anthropic account. If you don\'t have one:'
    )

    steps_2_3 = [
        'Open your browser and go to: https://console.anthropic.com/',
        'Click "Sign Up" and create an account',
        'Verify your email address',
        'Add a payment method or activate free trial credits'
    ]
    add_numbered_list(doc, steps_2_3)

    add_tip_box(doc, 'Claude CLI uses API credits. The cost is very low for typical use (a few dollars/month for moderate use).')

    # Step 2.4
    doc.add_heading('Step 2.4: Authenticate Claude CLI', level=2)

    doc.add_paragraph('In your Terminal, type:')
    add_code_block(doc, 'claude')

    doc.add_paragraph(
        'The first time you run this, Claude will ask you to authenticate. '
        'Follow the prompts—it will open a browser window where you\'ll log in to your Anthropic account '
        'and authorize Claude CLI.'
    )

    doc.add_paragraph('You may need to:')
    auth_steps = [
        'Press Enter when prompted',
        'Click a link that opens in your browser',
        'Log in to your Anthropic account',
        'Click "Authorize" to connect Claude CLI',
        'Return to VS Code—you should see Claude ready to chat'
    ]
    add_numbered_list(doc, auth_steps)

    add_success_box(doc, 'You see a prompt waiting for your input, something like "> " or a welcome message from Claude.')

    # Step 2.5
    doc.add_heading('Step 2.5: Your First Claude Conversation', level=2)

    doc.add_paragraph('Let\'s test everything with a simple project. First, exit Claude if it\'s running:')
    add_code_block(doc, '/exit')

    doc.add_paragraph('Create a test project folder and navigate into it:')
    add_code_block(doc, 'cd C:\\Users\\%USERNAME%\\Documents\\projects\nmkdir my-first-project\ncd my-first-project')

    doc.add_paragraph('Now start Claude in this folder:')
    add_code_block(doc, 'claude')

    doc.add_paragraph('Try your first vibe coding request. Type this and press Enter:')
    add_code_block(doc, 'Create a Python script called hello.py that prints "Hello! I just wrote my first vibe-coded program!"')

    doc.add_paragraph('Watch what happens:')
    what_happens = [
        'Claude thinks about your request',
        'Claude creates a file called hello.py',
        'You\'ll see the file appear in VS Code\'s Explorer sidebar on the left',
        'Claude shows you the code it wrote',
        'Claude may offer to run it for you'
    ]
    add_numbered_list(doc, what_happens)

    doc.add_paragraph('If Claude doesn\'t run it automatically, you can run it yourself. Exit Claude first:')
    add_code_block(doc, '/exit')

    doc.add_paragraph('Then run your script:')
    add_code_block(doc, 'python hello.py')

    add_success_box(doc, 'You see "Hello! I just wrote my first vibe-coded program!" printed in the Terminal.')

    p = doc.add_paragraph()
    run = p.add_run('Congratulations—you just vibe coded! ')
    run.bold = True
    p.add_run('You described what you wanted in English, and AI wrote the code.')

    # Phase 2 Checkpoint
    doc.add_heading('Phase 2 Checkpoint', level=2)

    checklist_2 = [
        '☐ Typing "claude" in Terminal starts the Claude interface',
        '☐ You can give Claude a simple instruction and it creates a file',
        '☐ You can see the new file in VS Code\'s Explorer sidebar',
        '☐ You can run a Python script with "python filename.py"',
        '☐ You know how to exit Claude (/exit)'
    ]
    for item in checklist_2:
        doc.add_paragraph(item)

    doc.add_page_break()

    # =========================================================================
    # PHASE 3
    # =========================================================================
    doc.add_heading('Phase 3: Project Structure & Best Practices', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Time estimate: 20-30 minutes')
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('Goal: ')
    run.bold = True
    p.add_run('Learn the professional way to organize projects so Claude works better and your code stays organized.')

    doc.add_paragraph(
        'This phase teaches you patterns that professionals use. They\'re not strictly required, '
        'but they\'ll prevent headaches as your projects grow.'
    )

    # Step 3.1
    doc.add_heading('Step 3.1: Understanding Virtual Environments', level=2)

    doc.add_paragraph(
        'A "virtual environment" (venv) is like a sandbox for your Python project. '
        'Here\'s why it matters:'
    )

    venv_reasons = [
        'Each project can have different versions of the same tools',
        'One project can\'t accidentally break another',
        'When you share your project, others can install exactly what they need'
    ]
    for item in venv_reasons:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('Think of it like having separate toolboxes for different jobs.')

    doc.add_paragraph('In your project folder, create a virtual environment:')
    add_code_block(doc, 'cd C:\\Users\\%USERNAME%\\Documents\\projects\\my-first-project\npython -m venv venv')

    doc.add_paragraph('This creates a "venv" folder inside your project. Now "activate" it:')
    add_code_block(doc, 'venv\\Scripts\\activate')

    add_success_box(doc, 'You see "(venv)" appear at the start of your Terminal prompt. This means you\'re working inside the virtual environment.')

    add_screenshot_placeholder(doc, "Terminal showing (venv) at the start of the prompt")

    add_warning_box(doc, 'Every time you open VS Code to work on a project, remember to activate the venv first!')

    doc.add_paragraph('To exit the virtual environment later (when you\'re done working):')
    add_code_block(doc, 'deactivate')

    # Step 3.2
    doc.add_heading('Step 3.2: Creating the .claude Folder', level=2)

    doc.add_paragraph(
        'The .claude folder contains configuration files that give Claude context about your project. '
        'These files help Claude understand your preferences and write better code for you.'
    )

    doc.add_paragraph('Create the folder:')
    add_code_block(doc, 'mkdir .claude')

    doc.add_heading('The Key Files', level=3)

    files_desc = [
        ('CLAUDE.md', 'The main instruction file. Claude reads this AUTOMATICALLY when you start a session. '
         'Put your project description, how to run things, and important context here. This is the most important file.'),
        ('AI_CODING_GUIDELINES.md', 'Rules for how Claude should write code. This prevents Claude from '
         'making unnecessary changes, "improving" things you didn\'t ask for, or going overboard.'),
        ('SESSION_LOG.md', 'A running log of what you accomplished in each session. Helps you (and Claude) '
         'remember context between sessions—like notes to your future self.')
    ]

    for name, desc in files_desc:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    # Step 3.3
    doc.add_heading('Step 3.3: Create Your CLAUDE.md File', level=2)

    doc.add_paragraph('Here\'s the magic of vibe coding—ask Claude to create these files for you!')

    doc.add_paragraph('Start Claude:')
    add_code_block(doc, 'claude')

    doc.add_paragraph('Then tell Claude what you need:')

    add_code_block(doc, '''Create a .claude/CLAUDE.md file for this project with:
- Project name: My First Project
- Description: A learning project for practicing vibe coding
- Virtual environment: venv/ (activate with "venv\\Scripts\\activate")
- A "Session Start" section that reminds me to read AI_CODING_GUIDELINES.md''')

    doc.add_paragraph('Claude will create a nicely formatted file. You can edit it later as your project evolves.')

    # Step 3.4
    doc.add_heading('Step 3.4: Add AI Coding Guidelines', level=2)

    doc.add_paragraph('These guidelines tell Claude how to behave. Without them, Claude might:')

    without_guidelines = [
        'Refactor code you didn\'t ask it to change',
        'Add features you didn\'t request',
        '"Improve" working code in ways that break it',
        'Make changes across multiple files when you asked for one small fix'
    ]
    for item in without_guidelines:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('Ask Claude to create guidelines:')
    add_code_block(doc, 'Create a .claude/AI_CODING_GUIDELINES.md file with these rules:\n1. State your approach before coding\n2. Only modify what I ask you to modify\n3. Don\'t "improve" or refactor code I didn\'t mention\n4. Ask if you\'re unsure rather than guessing')

    doc.add_paragraph('See the Appendix at the end of this guide for a complete template you can use.')

    # Step 3.5
    doc.add_heading('Step 3.5: Create a Session Log', level=2)

    doc.add_paragraph('The session log helps you pick up where you left off:')
    add_code_block(doc, 'Create a .claude/SESSION_LOG.md file with a template for logging sessions')

    # Phase 3 Checkpoint
    doc.add_heading('Phase 3 Checkpoint', level=2)

    doc.add_paragraph('Your project should now have this structure:')
    add_code_block(doc, '''my-first-project\\
├── .claude\\
│   ├── CLAUDE.md           ← Claude reads this automatically
│   ├── AI_CODING_GUIDELINES.md
│   └── SESSION_LOG.md
├── venv\\                   ← Your virtual environment (don't edit files in here)
└── hello.py                ← Your first script''')

    doc.add_paragraph('You can see this structure in VS Code\'s Explorer sidebar on the left.')

    doc.add_page_break()

    # =========================================================================
    # PHASE 4
    # =========================================================================
    doc.add_heading('Phase 4: GitHub Integration', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Time estimate: 20-30 minutes')
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('Goal: ')
    run.bold = True
    p.add_run('Connect your project to GitHub for backup and version history.')

    doc.add_paragraph(
        'GitHub is your safety net. Every change you save (called a "commit") is stored forever. '
        'If you break something, you can always go back to a working version.'
    )

    # Step 4.1
    doc.add_heading('Step 4.1: Create a GitHub Account', level=2)

    doc.add_paragraph('If you don\'t have a GitHub account:')
    steps_4_1 = [
        'Go to https://github.com/',
        'Click "Sign up"',
        'Choose a username (this will be visible publicly)',
        'Use the same email you configured Git with in Phase 1',
        'Verify your email address'
    ]
    add_numbered_list(doc, steps_4_1)

    # Step 4.2
    doc.add_heading('Step 4.2: Install GitHub CLI', level=2)

    doc.add_paragraph('The GitHub CLI (called "gh") makes it easy to work with GitHub from Terminal.')

    steps_gh = [
        'Go to: https://cli.github.com/',
        'Click "Download for Windows"',
        'Run the installer and accept defaults',
        'Close and reopen VS Code'
    ]
    add_numbered_list(doc, steps_gh)

    doc.add_paragraph('Now authenticate with your GitHub account:')
    add_code_block(doc, 'gh auth login')

    doc.add_paragraph('Follow the prompts:')
    gh_prompts = [
        'Choose "GitHub.com" (not Enterprise)',
        'Choose "HTTPS"',
        'When asked to authenticate, choose "Login with a web browser"',
        'Copy the one-time code shown',
        'Press Enter to open your browser',
        'Paste the code and click "Authorize"'
    ]
    add_numbered_list(doc, gh_prompts)

    add_success_box(doc, 'You see "Logged in as your-username"')

    # Step 4.3
    doc.add_heading('Step 4.3: Initialize Git in Your Project', level=2)

    doc.add_paragraph('Make sure you\'re in your project folder:')
    add_code_block(doc, 'cd C:\\Users\\%USERNAME%\\Documents\\projects\\my-first-project')

    doc.add_paragraph('Tell Git to start tracking this folder:')
    add_code_block(doc, 'git init')

    add_success_box(doc, 'You see "Initialized empty Git repository"')

    # Step 4.4
    doc.add_heading('Step 4.4: Create a .gitignore File', level=2)

    doc.add_paragraph('Some files shouldn\'t be uploaded to GitHub:')
    ignore_reasons = [
        'The venv/ folder (too big, can be recreated)',
        '.env files (may contain passwords or API keys)',
        'Temporary files'
    ]
    for item in ignore_reasons:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('A .gitignore file tells Git to skip these. Ask Claude:')
    add_code_block(doc, 'claude')

    doc.add_paragraph('Then:')
    add_code_block(doc, 'Create a .gitignore file for a Python project. Ignore: venv/, .env, __pycache__/, *.pyc')

    # Step 4.5
    doc.add_heading('Step 4.5: Make Your First Commit', level=2)

    doc.add_paragraph('Exit Claude first:')
    add_code_block(doc, '/exit')

    doc.add_paragraph('A "commit" is like a save point in a video game. It captures your code at this exact moment.')

    doc.add_paragraph('First, "stage" all your files (prepare them to be saved):')
    add_code_block(doc, 'git add .')

    doc.add_paragraph('Now create the commit with a message describing what you did:')
    add_code_block(doc, 'git commit -m "Initial project setup with Claude configuration"')

    add_success_box(doc, 'You see a message showing files were committed.')

    # Step 4.6
    doc.add_heading('Step 4.6: Push to GitHub', level=2)

    doc.add_paragraph('Create a new repository on GitHub and upload your code in one command:')
    add_code_block(doc, 'gh repo create my-first-project --private --source=. --push')

    doc.add_paragraph('This creates a PRIVATE repository (only you can see it) and uploads everything.')

    add_tip_box(doc, 'Use --private for personal projects. Change to --public only if you want to share.')

    # Phase 4 Checkpoint
    doc.add_heading('Phase 4 Checkpoint', level=2)

    checklist_4 = [
        '☐ Go to https://github.com/ and sign in',
        '☐ You see "my-first-project" in your repositories',
        '☐ Click on it—you see your files including the .claude folder',
        '☐ You understand: git add → git commit → git push'
    ]
    for item in checklist_4:
        doc.add_paragraph(item)

    p = doc.add_paragraph()
    run = p.add_run('You\'re fully set up! ')
    run.bold = True
    p.add_run('The next sections cover your daily workflow and troubleshooting.')

    doc.add_page_break()

    # =========================================================================
    # DAILY WORKFLOW
    # =========================================================================
    doc.add_heading('Your Daily Vibe Coding Workflow', level=1)

    doc.add_paragraph('Now that everything is set up, here\'s what a typical session looks like.')

    doc.add_heading('Starting a Session (2 minutes)', level=2)

    workflow_start = [
        'Open VS Code',
        'Open your project folder: File → Open Folder → Documents → projects → your-project',
        'Open Terminal: Ctrl+`',
        'Activate virtual environment: venv\\Scripts\\activate',
        'Start Claude: claude',
        'Tell Claude what you want to work on: "Today I want to..."'
    ]
    add_numbered_list(doc, workflow_start)

    add_tip_box(doc, 'Claude automatically reads your .claude/CLAUDE.md file, so it already knows about your project!')

    doc.add_heading('During Your Session', level=2)

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run('Good habits:')
    run.bold = True

    during_session = [
        'Describe what you want in plain English—be specific about the end result',
        'Review what Claude creates before saying "looks good"',
        'Test the code to make sure it actually works',
        'Give clear feedback: "That\'s not quite right, I wanted X instead" or "Perfect!"',
        'Ask Claude to explain anything you don\'t understand: "Explain what that code does"',
        'Take breaks—vibe coding is mentally engaging'
    ]
    for item in during_session:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Ending a Session (3 minutes)', level=2)

    workflow_end = [
        'Tell Claude: "Update the SESSION_LOG with what we did today"',
        'Exit Claude: /exit',
        'Save to GitHub:'
    ]
    add_numbered_list(doc, workflow_end)

    add_code_block(doc, 'git add .\ngit commit -m "Brief description of what you did"\ngit push')

    add_tip_box(doc, 'Commit often! Small, frequent commits are better than one huge commit at the end.')

    doc.add_page_break()

    # =========================================================================
    # TALKING TO CLAUDE
    # =========================================================================
    doc.add_heading('How to Talk to Claude Effectively', level=1)

    doc.add_paragraph(
        'The better you communicate with Claude, the better results you\'ll get. '
        'Here are patterns that work well.'
    )

    doc.add_heading('Be Specific About the End Result', level=2)

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run('Instead of:')
    run.italic = True
    doc.add_paragraph('"Analyze this data"', style='List Bullet')

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run('Say:')
    run.italic = True
    doc.add_paragraph('"Read the sales.xlsx file and create a summary showing total sales by region, with a chart"', style='List Bullet')

    doc.add_heading('Mention File Names', level=2)

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run('Instead of:')
    run.italic = True
    doc.add_paragraph('"Update the script to also show averages"', style='List Bullet')

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run('Say:')
    run.italic = True
    doc.add_paragraph('"Update analyze.py to also calculate and display the average for each column"', style='List Bullet')

    doc.add_heading('Share Error Messages', level=2)

    doc.add_paragraph(
        'When something breaks, copy the entire error message and paste it to Claude. '
        'Don\'t paraphrase—the exact text helps Claude diagnose the problem.'
    )

    doc.add_heading('Ask for Explanations', level=2)

    explain_examples = [
        '"Explain what this code does in simple terms"',
        '"Why did you use a list here instead of a dictionary?"',
        '"What would happen if the file doesn\'t exist?"'
    ]
    for item in explain_examples:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Give Clear Feedback', level=2)

    feedback_examples = [
        ('If it worked:', '"Perfect, that\'s exactly what I wanted"'),
        ('If it\'s close:', '"Almost, but change the date format to MM/DD/YYYY"'),
        ('If it\'s wrong:', '"That\'s not right. I wanted X, but this does Y. Try again."'),
        ('If you\'re confused:', '"I don\'t understand what this does. Can you explain?"')
    ]

    for situation, response in feedback_examples:
        p = doc.add_paragraph()
        run = p.add_run(f"{situation} ")
        run.bold = True
        p.add_run(response)

    doc.add_page_break()

    # =========================================================================
    # TROUBLESHOOTING
    # =========================================================================
    doc.add_heading('Troubleshooting Common Issues', level=1)

    doc.add_heading('"command not recognized" or "not found"', level=2)
    doc.add_paragraph('The tool isn\'t installed or isn\'t in your PATH.')
    fixes_1 = [
        'Close and reopen VS Code (this reloads the PATH)',
        'Make sure you checked "Add to PATH" during installation',
        'Try reinstalling the tool'
    ]
    for item in fixes_1:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('"Access denied" or "Permission denied"', level=2)
    doc.add_paragraph('You don\'t have admin rights or something is locked.')
    fixes_2 = [
        'Try running VS Code as Administrator (right-click → Run as administrator)',
        'Close any other programs that might have the file open',
        'Restart your computer and try again'
    ]
    for item in fixes_2:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Terminal seems frozen or stuck', level=2)
    fixes_3 = [
        'Press Ctrl+C to cancel the current operation',
        'If that doesn\'t work, close the Terminal panel and open a new one (Ctrl+`)',
        'As a last resort, close VS Code and reopen it'
    ]
    for item in fixes_3:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Claude made a mistake', level=2)
    doc.add_paragraph('This is normal! AI isn\'t perfect. Just tell Claude:')
    fixes_4 = [
        '"That\'s not right. Undo that change."',
        '"That broke the code. Here\'s the error: [paste error]. Fix it."',
        '"Go back to what we had before and try a different approach."'
    ]
    for item in fixes_4:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Can\'t find the Terminal', level=2)
    doc.add_paragraph('Press Ctrl+` (backtick key, left of the 1 key). Or go to View → Terminal.')

    doc.add_heading('Forgot to activate virtual environment', level=2)
    doc.add_paragraph('If you see errors about missing packages, you probably forgot to activate. Run:')
    add_code_block(doc, 'venv\\Scripts\\activate')
    doc.add_paragraph('You should see (venv) appear at the start of your prompt.')

    doc.add_heading('Git says "nothing to commit"', level=2)
    doc.add_paragraph('Either you haven\'t made any changes, or you already committed them. Check with:')
    add_code_block(doc, 'git status')

    doc.add_heading('PowerShell script execution is disabled', level=2)
    doc.add_paragraph('If you see an error about scripts being disabled when activating venv, run this once:')
    add_code_block(doc, 'Set-ExecutionPolicy -ExecutionPolicy RemoteSigned -Scope CurrentUser')
    doc.add_paragraph('Then try activating the venv again.')

    doc.add_page_break()

    # =========================================================================
    # TERMINAL COMMAND REFERENCE
    # =========================================================================
    doc.add_heading('Quick Reference: Terminal Commands', level=1)

    doc.add_paragraph('Keep this page handy until these become second nature.')

    doc.add_heading('Navigation', level=2)
    nav_commands = [
        ('cd', 'Print current folder path (where am I?)'),
        ('dir', 'List files in current folder'),
        ('cd foldername', 'Change to a folder'),
        ('cd ..', 'Go up one folder'),
        ('cd %USERPROFILE%', 'Go to your home folder'),
        ('cd %USERPROFILE%\\Documents\\projects', 'Go to your projects folder')
    ]
    for cmd, desc in nav_commands:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run.bold = True
        p.add_run(f"  —  {desc}")

    doc.add_heading('Python & Virtual Environment', level=2)
    python_commands = [
        ('python --version', 'Check Python version'),
        ('python -m venv venv', 'Create virtual environment'),
        ('venv\\Scripts\\activate', 'Activate virtual environment'),
        ('deactivate', 'Exit virtual environment'),
        ('python script.py', 'Run a Python script'),
        ('pip install package-name', 'Install a Python package (while venv active)')
    ]
    for cmd, desc in python_commands:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run.bold = True
        p.add_run(f"  —  {desc}")

    doc.add_heading('Claude CLI', level=2)
    claude_commands = [
        ('claude', 'Start Claude in current folder'),
        ('/exit', 'Exit Claude (type this while in Claude)'),
        ('/help', 'Show Claude help (while in Claude)')
    ]
    for cmd, desc in claude_commands:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run.bold = True
        p.add_run(f"  —  {desc}")

    doc.add_heading('Git (Version Control)', level=2)
    git_commands = [
        ('git status', 'See what files have changed'),
        ('git add .', 'Stage all changes for commit'),
        ('git commit -m "message"', 'Save a snapshot with description'),
        ('git push', 'Upload commits to GitHub'),
        ('git log --oneline', 'See recent commit history')
    ]
    for cmd, desc in git_commands:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run.bold = True
        p.add_run(f"  —  {desc}")

    doc.add_heading('Keyboard Shortcuts', level=2)
    shortcuts = [
        ('Ctrl+`', 'Toggle Terminal open/closed'),
        ('Ctrl+S', 'Save current file'),
        ('Ctrl+Shift+P', 'Open VS Code command palette'),
        ('Ctrl+C', 'Cancel current Terminal operation'),
        ('Up Arrow', 'Recall previous command')
    ]
    for cmd, desc in shortcuts:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run.bold = True
        p.add_run(f"  —  {desc}")

    doc.add_page_break()

    # =========================================================================
    # GLOSSARY
    # =========================================================================
    doc.add_heading('Glossary: Key Terms & Acronyms', level=1)

    doc.add_paragraph('Don\'t worry about memorizing these. Refer back as needed.')

    glossary = [
        ('API', 'Application Programming Interface. A way for programs to talk to each other. When Claude uses an API, it\'s communicating with Anthropic\'s servers.'),
        ('CLI', 'Command Line Interface. A text-based way to interact with software by typing commands (like Claude CLI).'),
        ('Commit', 'A saved snapshot of your code at a specific moment. Like a save point in a video game.'),
        ('Git', 'Version control software that tracks changes to your code over time.'),
        ('GitHub', 'A cloud service that stores Git repositories. Your code\'s online home.'),
        ('Node.js', 'A platform for running JavaScript code. Claude CLI is built on it.'),
        ('npm', 'Node Package Manager. Used to install Node.js tools like Claude CLI.'),
        ('PATH', 'A list of folders where Windows looks for programs. If a tool "isn\'t recognized," it\'s usually not in PATH.'),
        ('PowerShell', 'Windows\' modern command-line terminal. What you see when you open Terminal in VS Code.'),
        ('Push', 'Upload your commits from your computer to GitHub.'),
        ('Repository (Repo)', 'A project folder tracked by Git. Contains your code and its history.'),
        ('Terminal', 'A text-based interface for typing commands to your computer.'),
        ('Venv', 'Virtual Environment. An isolated Python workspace for a single project.'),
        ('VS Code', 'Visual Studio Code. A free code editor from Microsoft.')
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(definition)

    doc.add_page_break()

    # =========================================================================
    # KEY WEBSITES
    # =========================================================================
    doc.add_heading('Key Websites & Resources', level=1)

    doc.add_heading('Essential Sites', level=2)

    sites = [
        ('VS Code Download', 'https://code.visualstudio.com/', 'Download and install VS Code'),
        ('Python Download', 'https://www.python.org/downloads/', 'Download Python for Windows'),
        ('Git Download', 'https://git-scm.com/download/win', 'Download Git for Windows'),
        ('Node.js Download', 'https://nodejs.org/', 'Download Node.js'),
        ('GitHub CLI Download', 'https://cli.github.com/', 'Download GitHub CLI'),
        ('Anthropic Console', 'https://console.anthropic.com/', 'Create account, manage billing, get API keys'),
        ('GitHub', 'https://github.com/', 'Create account, view your repositories'),
        ('Claude CLI Docs', 'https://docs.anthropic.com/en/docs/claude-code', 'Official Claude CLI documentation')
    ]

    for name, url, desc in sites:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}")
        run.bold = True
        p.add_run(f"\n{url}\n{desc}")

    doc.add_heading('Learning Resources', level=2)

    learning = [
        ('Python Basics (if curious)', 'https://www.learnpython.org/', 'Free interactive Python tutorial'),
        ('Git Basics', 'https://git-scm.com/book/en/v2/Getting-Started-Git-Basics', 'Official Git documentation'),
        ('VS Code Tips', 'https://code.visualstudio.com/docs/getstarted/tips-and-tricks', 'VS Code productivity tips')
    ]

    for name, url, desc in learning:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}")
        run.bold = True
        p.add_run(f"\n{url}\n{desc}")

    doc.add_page_break()

    # =========================================================================
    # GETTING UNSTUCK
    # =========================================================================
    doc.add_heading('Getting Unstuck: When Things Go Wrong', level=1)

    doc.add_paragraph('Everyone gets stuck. Here\'s a systematic approach to getting unstuck.')

    doc.add_heading('Step 1: Don\'t Panic', level=2)
    doc.add_paragraph(
        'Getting stuck is normal. Even experienced developers spend significant time debugging. '
        'The goal isn\'t to never get stuck—it\'s to get unstuck efficiently.'
    )

    doc.add_heading('Step 2: Read the Error Message', level=2)
    doc.add_paragraph('Error messages often tell you exactly what\'s wrong. Look for:')
    error_parts = [
        'The file name and line number where the error occurred',
        'The type of error (e.g., "FileNotFoundError", "SyntaxError")',
        'A description of what went wrong'
    ]
    for item in error_parts:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Step 3: Ask Claude', level=2)
    doc.add_paragraph('Copy the entire error message and paste it to Claude with context:')
    add_code_block(doc, '''I ran "python analyze.py" and got this error:
[paste the full error here]

What's wrong and how do I fix it?''')

    doc.add_heading('Step 4: Check the Basics', level=2)
    doc.add_paragraph('Before going deeper, verify:')
    basics_check = [
        'Are you in the right folder? (type cd to check)',
        'Is your virtual environment activated? (look for (venv) in prompt)',
        'Did you save the file? (Ctrl+S)',
        'Are you running the right file? (check the filename)'
    ]
    for item in basics_check:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Step 5: Start Fresh', level=2)
    doc.add_paragraph('If nothing works:')
    fresh_start = [
        'Close VS Code completely and reopen it',
        'Open a new Terminal (Ctrl+`)',
        'Navigate back to your project folder',
        'Reactivate your virtual environment',
        'Try again'
    ]
    for item in fresh_start:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Step 6: Undo with Git', level=2)
    doc.add_paragraph('If Claude made changes that broke everything and you can\'t fix it, Git lets you go back:')
    add_code_block(doc, '# See what changed\ngit status\ngit diff\n\n# Undo all uncommitted changes (WARNING: loses your work since last commit)\ngit checkout .')

    add_warning_box(doc, 'git checkout . throws away all changes since your last commit. Only use this as a last resort.')

    doc.add_heading('Step 7: Ask for Help', level=2)
    doc.add_paragraph('If you\'re still stuck after trying the above:')
    help_options = [
        'Take a break and come back with fresh eyes',
        'Describe the problem to someone else (even explaining it often reveals the solution)',
        'Search for the error message online',
        'Reach out to whoever gave you this guide'
    ]
    for item in help_options:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Common "I\'m Stuck" Scenarios', level=2)

    scenarios = [
        ('"Claude isn\'t responding"',
         'Check your internet connection. Try /exit and restart with "claude". Check your Anthropic account has credits.'),
        ('"The code runs but does the wrong thing"',
         'Tell Claude specifically what\'s wrong: "It\'s calculating X but I expected Y because..."'),
        ('"I don\'t understand what Claude did"',
         'Ask: "Explain what this code does step by step in simple terms"'),
        ('"I made a mess and want to start over"',
         'Create a new project folder and start fresh. Your old folder is still there if you need it.'),
        ('"Nothing works and I\'ve tried everything"',
         'Sometimes the best move is to close everything, take a walk, and come back tomorrow.')
    ]

    for scenario, solution in scenarios:
        p = doc.add_paragraph()
        run = p.add_run(f"{scenario}")
        run.bold = True
        doc.add_paragraph(solution)

    doc.add_page_break()

    # =========================================================================
    # APPENDIX: TEMPLATE FILES
    # =========================================================================
    doc.add_heading('Appendix: Template Files', level=1)

    doc.add_paragraph('Copy these templates into your .claude folder, or ask Claude to create them for you.')

    doc.add_heading('CLAUDE.md Template', level=2)

    add_code_block(doc, '''# Project Name

## Overview
Brief description of what this project does.

## Environment Setup
- **Virtual Environment:** venv\\
- **Activate:** venv\\Scripts\\activate

## How to Run
- Main script: python main.py
- (Add your specific run commands here)

## Important Files
- main.py — The main script
- (List your key files here)

## Session Start
Before starting work, read:
1. .claude/AI_CODING_GUIDELINES.md

## End of Session
Update .claude/SESSION_LOG.md with what was accomplished.''')

    doc.add_heading('AI_CODING_GUIDELINES.md Template', level=2)

    add_code_block(doc, '''# AI Coding Guidelines

## Core Rules

1. **State your approach first**
   Before making changes, briefly describe what you plan to do and wait for confirmation.

2. **Only change what I ask for**
   Don't refactor, "improve," or clean up code I didn't mention.
   If you notice other issues, mention them but don't fix them unless I ask.

3. **Read before writing**
   Always read existing code before modifying it.
   Don't guess at how things work.

4. **Verify your work**
   Test code when possible. If you can't test it, say so.

5. **Ask when unsure**
   If you're not certain about something, ask rather than guess.

## Working Style
- Make one change at a time
- Keep solutions simple
- Explain what you did after making changes''')

    doc.add_heading('SESSION_LOG.md Template', level=2)

    add_code_block(doc, '''# Session Log

Add new entries at the bottom. Keep each entry brief.

---

## 2024-XX-XX: First Session
- **What was done:** Set up project structure
- **Key decisions:** Using Python 3.x, created basic folder structure
- **Next steps:** Start building first feature

---

(Add new entries below this line)''')

    # Save
    doc.save('/Users/sjmoorad/Documents/scott_projects/Vibe_Coding_Setup_Guide_Windows.docx')
    print("Document created: Vibe_Coding_Setup_Guide_Windows.docx")


if __name__ == "__main__":
    create_document()
